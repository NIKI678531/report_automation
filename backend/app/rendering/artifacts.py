from __future__ import annotations

import hashlib
from html.parser import HTMLParser
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.storage import storage
from app.domain.document import render_content_manifest, review_display_title
from app.domain.models import RenderArtifact, Report, ReportDocument
from .html import pct, price, render_html


MIME = {"html": "text/html", "pdf": "application/pdf", "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _page_setup(section, page_number: int) -> None:
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin, section.right_margin, section.bottom_margin, section.left_margin = Cm(0.7), Cm(1), Cm(1.5), Cm(1)
    header = section.header.paragraphs[0]
    header.text = "Monthly Commentary"
    header.style = "Header"
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run().add_picture(str(Path(__file__).parent / "static" / "csop-logo.png"), width=Cm(4.4))
    footer.add_run(f"  {page_number}")


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _table(document: Document, headers: list[str], rows: list[list[str]], blue_first: bool = False):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        _shade(cell, "2660AD")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if blue_first and i == 0:
                _shade(cells[i], "2660AD")
                for run in cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)
    return table


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag in {"p", "li", "h2", "h3", "blockquote"}:
            self.parts.append("\n")


def _plain_html(value: str) -> str:
    parser = _TextExtractor()
    parser.feed(value)
    return " ".join("".join(parser.parts).split())


def render_docx(report: Report, content: dict, destination: Path) -> None:
    sections = content["sections"]
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.color.rgb = RGBColor(34, 50, 127)
    _page_setup(document.sections[0], 1)
    document.add_heading(report.product_name, 0)
    review = sections["month_in_review"]
    enable_review_layout = content.get("template_version") != "3033-v1"
    document.add_heading(review_display_title(content), 1)
    if enable_review_layout and review.get("blocks"):
        grouped: dict[int, list[dict]] = {}
        for block in sorted(review["blocks"], key=lambda item: (item["y"], item["x"], item["block_id"])):
            grouped.setdefault(block["y"], []).append(block)
        for block_row in grouped.values():
            layout_table = document.add_table(rows=1, cols=len(block_row))
            layout_table.autofit = False
            for cell, block in zip(layout_table.rows[0].cells, block_row):
                cell.width = Inches(7.1 * block["w"] / 12)
                heading = cell.paragraphs[0]
                run = heading.add_run(block["title"])
                run.bold = True
                run.font.color.rgb = RGBColor(34, 50, 127)
                cell.add_paragraph(_plain_html(block["content"]))
    else:
        document.add_paragraph(review["summary"])
        document.add_heading("Key Drivers of the Correction", 1)
        for item in review["drivers"]:
            document.add_paragraph(f"{item['title']}\n{item['body']}", style="List Number")
        document.add_heading("Key Areas to Monitor", 1)
        for item in review["monitor"]:
            document.add_paragraph(f"{item['title']}\n{item['body']}", style="List Number")
        document.add_heading("Outlook", 1)
        document.add_paragraph(review["outlook"])
    document.add_heading(f"Historical Performance of {content['product_ticker']} and {content['benchmark_name']}*", 1)
    history = sections["historical_performance"]["rows"]
    _table(document, ["", "1-month return (%)", "3-month return (%)", "6-month return (%)", "YTD return (%)"], [[x["name"], pct(x["return_1m"]), pct(x["return_3m"]), pct(x["return_6m"]), pct(x["return_ytd"])] for x in history])
    document.add_paragraph(sections["footnotes"].get("historical", ""), style="Caption")

    _page_setup(document.add_section(WD_SECTION.NEW_PAGE), 2)
    document.add_heading("Company News", 1)
    for item in sections["company_news"]:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item["title"])
        run.bold = True; run.font.color.rgb = RGBColor(38, 96, 173)
        paragraph.add_run("\n" + item["summary"])

    _page_setup(document.add_section(WD_SECTION.NEW_PAGE), 3)
    heading = document.add_heading(f"The Performance of {getattr(report, 'constituent_index_code', report.benchmark_code)} Constituents", 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"(*Next Rebalancing Date: {content.get('next_rebalancing_date') or 'N/A'})").alignment = WD_ALIGN_PARAGRAPH.CENTER
    constituents = sections["constituents"]
    _table(document, ["Stock Code", "Stock Name", "Closing Price (HKD)", "Weighting (%)", "1-month return (%)", "3-month return (%)", "6-month return (%)", "YTD return (%)"], [[x["security_code"], x["name_en"], price(x["close_price"]), pct(x["weight"]), pct(x["return_1m"]), pct(x["return_3m"]), pct(x["return_6m"]), pct(x["return_ytd"])] for x in constituents], blue_first=True)
    document.add_paragraph(sections["footnotes"].get("constituents", ""), style="Caption")

    _page_setup(document.add_section(WD_SECTION.NEW_PAGE), 4)
    analytics = sections["analytics"]
    document.add_heading("Top 10 Index Constituents* (%)", 1)
    _table(document, ["Issuer", "Weight (%)"], [[x["issuer"], pct(x["weight"])] for x in analytics["top10"]])
    document.add_heading("Index Sectors Breakdown*", 1)
    _table(document, ["Sector", "Weight (%)"], [[x["sector"], pct(x["weight"])] for x in analytics["sectors"]])
    document.add_heading(f"Top Performers in {content['month_name']}", 1)
    _table(document, ["Issuer", "Return (%)"], [[x["issuer"], pct(x["return"])] for x in analytics["top"]])
    document.add_heading(f"Bottom Performers in {content['month_name']}", 1)
    _table(document, ["Issuer", "Return (%)"], [[x["issuer"], pct(x["return"])] for x in analytics["bottom"]])
    document.add_heading(f"{content['product_ticker']} Portfolio Analysis", 1)
    _table(document, ["Measure", "Value"], [[x["label"], x["value"]] for x in analytics["portfolio"]])
    document.add_paragraph(sections["footnotes"].get("analytics", ""), style="Caption")
    document.save(destination)


def build_artifact(db: Session, report: Report, document: ReportDocument, format_name: str) -> RenderArtifact:
    if format_name not in MIME:
        raise ValueError(f"Unsupported format: {format_name}")
    directory = settings.output_root / format_name
    directory.mkdir(parents=True, exist_ok=True)
    destination = directory / f"{report.product_code}_{report.report_date.isoformat()}_v{document.version}.{format_name}"
    html = render_html(report, document.content)
    if format_name == "html":
        destination.write_text(html, encoding="utf-8")
    elif format_name == "docx":
        render_docx(report, document.content, destination)
    else:
        with TemporaryDirectory() as temp:
            source = Path(temp) / "report.html"
            source.write_text(html, encoding="utf-8")
            with sync_playwright() as playwright:
                browser = playwright.chromium.launch(headless=True)
                page = browser.new_page()
                page.goto(source.as_uri(), wait_until="networkidle")
                page.pdf(path=str(destination), format="A4", print_background=True, margin={"top": "0", "right": "0", "bottom": "0", "left": "0"}, prefer_css_page_size=True)
                browser.close()
    object_key = f"{format_name}/{destination.name}"
    stored = storage.put_file(destination, object_key)
    content_manifest = render_content_manifest(document.content)
    existing = list(db.scalars(select(RenderArtifact).where(
        RenderArtifact.report_id == report.id,
        RenderArtifact.document_version == document.version,
    )))
    if any(item.content_manifest.get("checksum") != content_manifest["checksum"] for item in existing):
        raise ValueError("QC-010: canonical content manifest differs across output formats")
    artifact = RenderArtifact(
        report_id=report.id,
        document_version=document.version,
        format=format_name,
        storage_key=stored.key,
        mime_type=MIME[format_name],
        size_bytes=stored.size_bytes,
        checksum=stored.checksum,
        template_version=document.template_version,
        renderer_version=settings.renderer_version if format_name == "pdf" else f"{format_name}-v1",
        content_manifest=content_manifest,
    )
    db.add(artifact)
    db.commit()
    db.refresh(artifact)
    return artifact
