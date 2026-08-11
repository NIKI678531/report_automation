import io
from datetime import datetime, timezone

from docx import Document


def prepared_report(client):
    report = client.post("/api/v1/reports", json={"report_date": "2026-06-30"}).json()
    client.post(f"/api/v1/reports/{report['id']}/snapshots", json={"source_policy": "GOLDEN_FIXTURE"})
    client.post(f"/api/v1/reports/{report['id']}/calculations")
    return report["id"]


def test_calculation_and_ai_draft_are_versioned_and_bound(client):
    report_id = prepared_report(client)
    calculated = client.post(f"/api/v1/reports/{report_id}/calculations")
    assert calculated.status_code == 200, calculated.text
    assert calculated.json()["metrics"]["constituent_count"] == 30
    assert calculated.json()["formula_version"] == "hstech-2026.1"
    assert client.get(f"/api/v1/reports/{report_id}").json()["status"] == "EDITING"
    version = calculated.json()["document_version"]
    metrics = client.get(f"/api/v1/reports/{report_id}/metrics").json()
    modules = client.get(f"/api/v1/reports/{report_id}/modules").json()
    quality = client.get(f"/api/v1/reports/{report_id}/quality-results").json()
    assert {item["metric_code"] for item in metrics} >= {
        "constituent_count", "weight_total", "historical.return_1m",
        "constituent.close_price", "constituent.weight", "constituent.return_1m", "industry.weight",
    }
    assert len([item for item in metrics if item["metric_code"] == "constituent.weight"]) == 30
    assert {item["module_code"] for item in modules} == {
        "constituents_performance", "final_analytics", "footnotes", "historical_performance",
    }
    assert all(item["source_dataset_ids"] for item in modules)
    assert {item["check_id"] for item in quality} >= {"QC-001", "QC-002", "QC-004"}
    bound = client.get(f"/api/v1/reports/{report_id}").json()["latest_document"]["content"]["module_bindings"]
    assert {value["module_snapshot_id"] for value in bound.values()} == {item["id"] for item in modules}
    drafted = client.post(f"/api/v1/reports/{report_id}/ai/in-review", json={"version": version, "user_prompt": "Approved outlook pending reviewer confirmation."})
    assert drafted.status_code == 200, drafted.text
    content = drafted.json()["content"]
    assert content["ai_provenance"]["provider"] == "deterministic-template"
    assert content["ai_provenance"]["metric_bindings"]
    assert "30 constituents" in content["sections"]["month_in_review"]["summary"]
    review = client.get(f"/api/v1/reports/{report_id}/review")
    assert next(item for item in review.json()["checks"] if item["check_id"] == "QC-008")["status"] == "PASSED"


def test_ai_number_check_blocks_unbound_numbers(client):
    report_id = prepared_report(client)
    detail = client.get(f"/api/v1/reports/{report_id}").json()
    drafted = client.post(
        f"/api/v1/reports/{report_id}/ai/in-review",
        json={"version": detail["latest_document"]["version"], "user_prompt": "Approved outlook."},
    ).json()
    content = drafted["content"]
    content["sections"]["month_in_review"]["outlook"] = "The fund is expected to return 99%."
    saved = client.patch(
        f"/api/v1/reports/{report_id}/document",
        json={"version": drafted["version"], "content": content},
    )
    assert saved.status_code == 200, saved.text
    review = client.get(f"/api/v1/reports/{report_id}/review").json()
    check = next(item for item in review["checks"] if item["check_id"] == "QC-008")
    assert check["status"] == "FAILED"
    assert "99%" in check["actual"]["unmatched"]
    finalized = client.post(
        f"/api/v1/reports/{report_id}/finalize",
        json={"version": saved.json()["version"]},
    )
    assert finalized.status_code == 422
    assert finalized.json()["error_code"] == "QC-008"
    assert client.get(f"/api/v1/reports/{report_id}").json()["status"] == "QA_BLOCKED"


def test_news_candidate_selection_and_order(client):
    report_id = prepared_report(client)
    created = []
    for index in range(2):
        response = client.post(f"/api/v1/reports/{report_id}/news/candidates", json={
            "source_name": "Approved source", "source_url": f"https://example.test/news-{index}",
            "published_at": datetime(2026, 6, index + 1, tzinfo=timezone.utc).isoformat(),
            "title": f"News {index}", "summary": f"Summary {index}", "security_code": "700", "ticker": "0700.HK", "importance": "HIGH",
        })
        assert response.status_code == 201
        created.append(response.json())
    detail = client.get(f"/api/v1/reports/{report_id}").json()
    version = detail["latest_document"]["version"]
    selected = client.put(f"/api/v1/reports/{report_id}/news", json={"version": version, "items": [
        {"news_item_id": created[1]["id"], "position": 0},
        {"news_item_id": created[0]["id"], "position": 1, "title_override": "Reviewed title"},
    ]})
    assert selected.status_code == 200, selected.text
    assert [item["title"] for item in selected.json()["items"]] == ["News 1", "Reviewed title"]
    preview = client.post(f"/api/v1/reports/{report_id}/preview")
    assert preview.text.index("News 1") < preview.text.index("Reviewed title")


def test_manual_news_must_be_inside_the_report_month(client):
    report_id = prepared_report(client)
    response = client.post(f"/api/v1/reports/{report_id}/news/candidates", json={
        "source_name": "Source",
        "source_url": "https://example.test/news-before-month",
        "published_at": datetime(2026, 5, 31, tzinfo=timezone.utc).isoformat(),
        "title": "Before the report month",
        "summary": "Out of range",
        "ticker": "0700.HK",
    })

    assert response.status_code == 422, response.text
    assert response.json()["error_code"] == "NEWS_DATE_RANGE_INVALID"


def test_selected_news_survives_recalculation_and_renders_month_metadata(client):
    report_id = prepared_report(client)
    candidate = client.post(f"/api/v1/reports/{report_id}/news/candidates", json={
        "source_name": "Reuters",
        "source_url": "https://example.test/june-news",
        "published_at": datetime(2026, 6, 12, 8, 30, tzinfo=timezone.utc).isoformat(),
        "title": "June selected headline",
        "summary": "Selected report-month summary",
        "ticker": "0700.HK",
    })
    assert candidate.status_code == 201, candidate.text
    detail = client.get(f"/api/v1/reports/{report_id}").json()
    selected = client.put(f"/api/v1/reports/{report_id}/news", json={
        "version": detail["latest_document"]["version"],
        "items": [{"news_item_id": candidate.json()["id"], "position": 0}],
    })
    assert selected.status_code == 200, selected.text

    recalculated = client.post(f"/api/v1/reports/{report_id}/calculations")

    assert recalculated.status_code == 200, recalculated.text
    content = client.get(f"/api/v1/reports/{report_id}").json()["latest_document"]["content"]
    assert [item["title"] for item in content["sections"]["company_news"]] == ["June selected headline"]
    preview = client.post(f"/api/v1/reports/{report_id}/preview")
    assert "Reuters" in preview.text
    assert "2026-06-12" in preview.text
    assert "https://example.test/june-news" in preview.text


def test_review_accepts_complete_golden_editorial(client):
    report_id = prepared_report(client)
    review = client.get(f"/api/v1/reports/{report_id}/review")
    assert review.status_code == 200
    assert review.json()["ready"] is True
    assert any(item["check_id"] == "QC-009" and item["status"] == "PASSED" for item in review.json()["checks"])


def test_review_layout_is_sanitized_versioned_and_rejects_overlap(client):
    report_id = prepared_report(client)
    detail = client.get(f"/api/v1/reports/{report_id}").json()
    content = detail["latest_document"]["content"]
    version = detail["latest_document"]["version"]
    content["sections"]["month_in_review"]["display_title"] = "June Market Reset"
    content["sections"]["month_in_review"]["title"] = "June Market Reset"
    content["sections"]["month_in_review"]["blocks"] = [
        {"block_id": "summary", "type": "rich_text", "title": "Summary", "content": '<p>Approved</p><script>alert(1)</script><a href="javascript:bad">bad</a>', "x": 0, "y": 0, "w": 12, "h": 4},
        {"block_id": "outlook", "type": "outlook", "title": "Outlook", "content": "<p>Outlook</p>", "x": 0, "y": 4, "w": 6, "h": 4},
    ]
    saved = client.patch(f"/api/v1/reports/{report_id}/document", json={"version": version, "content": content})
    assert saved.status_code == 200, saved.text
    review = saved.json()["content"]["sections"]["month_in_review"]
    assert review["title"] == "June Market Reset"
    assert review["display_title"] == "June Market Reset"
    assert "<script>" not in review["blocks"][0]["content"]
    assert "javascript:" not in review["blocks"][0]["content"]
    preview = client.post(f"/api/v1/reports/{report_id}/preview")
    assert preview.status_code == 200
    assert ">June Market Reset</h2>" in preview.text
    assert 'data-block-id="summary"' in preview.text
    assert "<script>" not in preview.text

    invalid_content = saved.json()["content"]
    invalid_content["sections"]["month_in_review"]["blocks"][1].update({"x": 5, "y": 2})
    rejected = client.patch(f"/api/v1/reports/{report_id}/document", json={"version": saved.json()["version"], "content": invalid_content})
    assert rejected.status_code == 422
    assert rejected.json()["error_code"] == "REVIEW_LAYOUT_INVALID"


def test_review_blocks_render_to_html_and_editable_docx(client):
    report_id = prepared_report(client)
    detail = client.get(f"/api/v1/reports/{report_id}").json()
    content = detail["latest_document"]["content"]
    content["sections"]["month_in_review"]["display_title"] = "Custom June Review"
    content["sections"]["month_in_review"]["title"] = "Custom June Review"
    content["sections"]["month_in_review"]["blocks"] = [
        {"block_id": "summary", "type": "rich_text", "title": "Summary", "content": "<p>Editable custom review content</p>", "x": 0, "y": 0, "w": 12, "h": 4},
        {"block_id": "outlook", "type": "outlook", "title": "Outlook", "content": "<p>Approved outlook</p>", "x": 0, "y": 4, "w": 6, "h": 4},
    ]
    saved = client.patch(f"/api/v1/reports/{report_id}/document", json={"version": detail["latest_document"]["version"], "content": content})
    assert saved.status_code == 200, saved.text
    finalized = client.post(f"/api/v1/reports/{report_id}/finalize", json={"version": saved.json()["version"]})
    assert finalized.status_code == 200, finalized.text
    rendered = client.post(
        f"/api/v1/reports/{report_id}/renders",
        json={"formats": ["html", "docx"]},
        headers={"Idempotency-Key": f"review-layout-{report_id}"},
    )
    assert rendered.status_code == 202, rendered.text
    assert [job["status"] for job in rendered.json()] == ["SUCCEEDED", "SUCCEEDED"]

    artifacts = {}
    for job in rendered.json():
        signed = client.get(f"/api/v1/artifacts/{job['artifact_id']}/download").json()
        artifacts[job["format"]] = client.get(signed["download_url"]).content
    assert b'data-block-id="summary"' in artifacts["html"]
    assert b"Editable custom review content" in artifacts["html"]
    docx = Document(io.BytesIO(artifacts["docx"]))
    text = "\n".join(
        [paragraph.text for paragraph in docx.paragraphs]
        + [cell.text for table in docx.tables for row in table.rows for cell in row.cells]
    )
    assert "Custom June Review" in text
    assert "Editable custom review content" in text


def test_review_title_survives_snapshot_rebinding_and_has_structured_validation(client):
    report = client.post("/api/v1/reports", json={"report_date": "2026-06-30"}).json()
    detail = client.get(f"/api/v1/reports/{report['id']}").json()
    content = detail["latest_document"]["content"]
    assert content["sections"]["month_in_review"]["display_title"] == "June in Review"
    content["sections"]["month_in_review"]["display_title"] = "Editable Monthly Perspective"
    content["sections"]["month_in_review"]["title"] = "Editable Monthly Perspective"
    saved = client.patch(
        f"/api/v1/reports/{report['id']}/document",
        json={"version": detail["latest_document"]["version"], "content": content},
    )
    assert saved.status_code == 200, saved.text

    rebound = client.post(f"/api/v1/reports/{report['id']}/snapshots", json={"source_policy": "GOLDEN_FIXTURE"})
    assert rebound.status_code == 201, rebound.text
    refreshed = client.get(f"/api/v1/reports/{report['id']}").json()["latest_document"]["content"]
    assert refreshed["sections"]["month_in_review"]["display_title"] == "Editable Monthly Perspective"
    assert ">Editable Monthly Perspective</h2>" in client.post(f"/api/v1/reports/{report['id']}/preview").text

    refreshed["sections"]["month_in_review"]["display_title"] = "   "
    rejected = client.patch(
        f"/api/v1/reports/{report['id']}/document",
        json={"version": 3, "content": refreshed},
    )
    assert rejected.status_code == 422
    assert rejected.json() == {
        "error_code": "REVIEW_TITLE_INVALID",
        "field": "sections.month_in_review.display_title",
        "entity_id": None,
        "message": "Review title cannot be empty.",
        "severity": "BLOCKING",
        "fix_hint": "Enter a Review title before saving.",
        "request_id": None,
    }
