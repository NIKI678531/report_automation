from types import SimpleNamespace

from docx import Document

from app.rendering.artifacts import render_docx


def test_products_are_effective_dated_and_drive_report_identity(client):
    products = client.get("/api/v1/products", params={"as_of_date": "2026-06-30"})
    assert products.status_code == 200
    assert ("3033", "3033.HK") in [(item["product_code"], item["ticker"]) for item in products.json()]
    assert ("3037", "3037.HK") in [(item["product_code"], item["ticker"]) for item in products.json()]

    created = client.post("/api/v1/reports", json={"product_code": "3033", "report_date": "2026-06-30"})
    assert created.status_code == 201, created.text
    assert created.json()["product_name"] == "CSOP Hang Seng TECH Index ETF (3033.HK)"
    assert created.json()["benchmark_code"] == "HSTECH"
    assert created.json()["template_version"] == "3033-v2"


def test_report_identity_changes_with_selected_etf(client):
    created = client.post("/api/v1/reports", json={"product_code": "3037", "report_date": "2026-06-30"})
    assert created.status_code == 201, created.text
    report = created.json()
    assert report["product_name"] == "CSOP Hang Seng Index ETF (3037.HK)"
    assert report["benchmark_code"] == "HSI"

    detail = client.get(f"/api/v1/reports/{report['id']}")
    assert detail.status_code == 200, detail.text
    content = detail.json()["latest_document"]["content"]
    assert content["product_ticker"] == "3037.HK"
    assert content["benchmark_name"] == "Hang Seng Index"

    preview = client.post(f"/api/v1/reports/{report['id']}/preview")
    assert preview.status_code == 200, preview.text
    assert '<div class="sector-empty">N/A</div>' in preview.text
    assert 'class="donut"' not in preview.text


def test_report_creation_rejects_unknown_or_client_supplied_product_identity(client):
    missing = client.post("/api/v1/reports", json={"product_code": "9999", "report_date": "2026-06-30"})
    assert missing.status_code == 422
    assert missing.json()["error_code"] == "PRODUCT_NOT_AVAILABLE"

    spoofed = client.post("/api/v1/reports", json={
        "product_code": "3033",
        "product_name": "Spoofed fund",
        "benchmark_code": "SPOOF",
        "report_date": "2026-06-30",
    })
    assert spoofed.status_code == 422


def test_non_3033_product_uses_its_own_count_formula_and_never_golden_data(client, tmp_path):
    report = client.post("/api/v1/reports", json={"product_code": "TEST", "report_date": "2026-06-30"})
    assert report.status_code == 201, report.text
    report_id = report.json()["id"]

    golden = client.post(f"/api/v1/reports/{report_id}/snapshots", json={"source_policy": "GOLDEN_FIXTURE"})
    assert golden.status_code == 422
    assert golden.json()["error_code"] == "FIXTURE_NOT_AVAILABLE"

    csv_data = (
        "security_code,ticker,name_en,close_price,currency,weight,sector,return_1m\n"
        "1,0001.HK,Alpha,10,HKD,0.5,Technology,0.1\n"
        "2,0002.HK,Beta,20,HKD,0.5,Financials,-0.1\n"
    )
    imported = client.post(
        f"/api/v1/reports/{report_id}/imports",
        files={"file": ("constituents.csv", csv_data, "text/csv")},
        data={"dataset_type": "constituents"},
    )
    assert imported.status_code == 201, imported.text
    assert all(item["status"] == "PASSED" for item in imported.json()["validation_results"])

    applied = client.post(
        f"/api/v1/reports/{report_id}/imports/{imported.json()['id']}/apply",
        json={"reason": "Approved synthetic test dataset"},
    )
    assert applied.status_code == 200, applied.text
    assert len(applied.json()["payload"]["constituents"]) == 2
    assert applied.json()["payload"]["historical_performance"] == {"rows": []}

    calculated = client.post(f"/api/v1/reports/{report_id}/calculations")
    assert calculated.status_code == 200, calculated.text
    assert calculated.json()["formula_version"] == "test-index-v1"
    assert calculated.json()["metrics"]["constituent_count"] == 2
    detail = client.get(f"/api/v1/reports/{report_id}").json()
    preview = client.post(f"/api/v1/reports/{report_id}/preview")
    assert preview.status_code == 200
    assert "Historical Performance of 9999.HK and Synthetic Test Index" in preview.text
    assert "The Performance of TESTIDX Constituents" in preview.text
    assert "9999.HK Portfolio Analysis" in preview.text
    assert "3033.HK" not in preview.text

    destination = tmp_path / "test-fund.docx"
    render_docx(
        SimpleNamespace(product_name=report.json()["product_name"], benchmark_code=report.json()["benchmark_code"]),
        detail["latest_document"]["content"],
        destination,
    )
    docx = Document(destination)
    text = "\n".join(
        [paragraph.text for paragraph in docx.paragraphs]
        + [cell.text for table in docx.tables for row in table.rows for cell in row.cells]
    )
    assert "Historical Performance of 9999.HK and Synthetic Test Index" in text
    assert "9999.HK Portfolio Analysis" in text
    assert "3033.HK" not in text


def test_document_update_rebinds_system_owned_report_identity(client):
    report = client.post("/api/v1/reports", json={"product_code": "TEST", "report_date": "2026-06-30"}).json()
    detail = client.get(f"/api/v1/reports/{report['id']}").json()
    content = detail["latest_document"]["content"]
    content.update({
        "report_id": "spoofed",
        "report_date": "2025-01-01",
        "month_name": "January",
        "product_ticker": "SPOOF.HK",
        "benchmark_name": "Spoofed Index",
        "template_version": "spoof-v1",
        "design_token_version": "spoof-v1",
        "language_mode": "SPOOF",
    })
    saved = client.put(
        f"/api/v1/reports/{report['id']}/document",
        json={"version": detail["latest_document"]["version"], "content": content},
    )
    assert saved.status_code == 200, saved.text
    identity = saved.json()["content"]
    assert identity["report_id"] == report["id"]
    assert identity["report_date"] == "2026-06-30"
    assert identity["month_name"] == "June"
    assert identity["product_ticker"] == "9999.HK"
    assert identity["benchmark_name"] == "Synthetic Test Index"
    assert identity["template_version"] == "test-v1"
    assert identity["design_token_version"] == "test-v1"
    assert identity["language_mode"] == "EN"


def test_admin_can_import_an_approved_product_catalog(client):
    catalog = (
        "product_code,ticker,name_en,name_zh_hant,benchmark_code,benchmark_name,currency,timezone,valid_from,valid_to,is_active,display_order,template_version,design_token_version,expected_constituent_count,formula_profile\n"
        "3067,3067.HK,CSOP Approved Test ETF,,APPROVEDIDX,Approved Index,HKD,Asia/Hong_Kong,2026-01-01,,true,20,monthly-v2,monthly-v2,50,approved-index-v1\n"
    )
    imported = client.post(
        "/api/v1/products/import",
        files={"file": ("approved-products.csv", catalog, "text/csv")},
        headers={"X-User-Role": "ADMIN", "X-Request-ID": "catalog-import"},
    )
    assert imported.status_code == 200, imported.text
    assert imported.json() == {"created": 1, "updated": 0, "total": 1}

    products = client.get("/api/v1/products", params={"as_of_date": "2026-06-30"}).json()
    assert any(item["product_code"] == "3067" and item["source"] == "APPROVED_IMPORT" for item in products)
    report = client.post("/api/v1/reports", json={"product_code": "3067", "report_date": "2026-06-30"})
    assert report.status_code == 201, report.text
    assert report.json()["product_name"] == "CSOP Approved Test ETF (3067.HK)"
    assert report.json()["benchmark_code"] == "APPROVEDIDX"


def test_product_catalog_import_is_admin_only_and_atomic(client):
    denied = client.post(
        "/api/v1/products/import",
        files={"file": ("approved-products.csv", b"invalid", "text/csv")},
        headers={"X-User-Role": "EDITOR"},
    )
    assert denied.status_code == 403
    assert denied.json()["error_code"] == "PRODUCT_ADMIN_REQUIRED"

    invalid = (
        "product_code,ticker,name_en,benchmark_code,valid_from,template_version,design_token_version,formula_profile\n"
        "3067,3067.HK,Fund A,IDX,2026-01-01,v1,v1,f1\n"
        "3067,3067.HK,Fund B,IDX,2026-01-01,v1,v1,f1\n"
    )
    rejected = client.post(
        "/api/v1/products/import",
        files={"file": ("approved-products.csv", invalid, "text/csv")},
        headers={"X-User-Role": "ADMIN"},
    )
    assert rejected.status_code == 422
    products = client.get("/api/v1/products", params={"as_of_date": "2026-06-30"}).json()
    assert not any(item["product_code"] == "3067" for item in products)
